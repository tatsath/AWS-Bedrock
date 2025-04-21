####################################################################
#                         import
####################################################################

import warnings

warnings.filterwarnings("ignore", category=FutureWarning)

import os, glob
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import html
from datetime import datetime

# Download NLTK data
import nltk
nltk.download('punkt')

# Import openai and google_genai as main LLM services
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_google_genai import GoogleGenerativeAIEmbeddings

# langchain prompts, memory, chains...
from langchain.prompts import PromptTemplate, ChatPromptTemplate
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory, ConversationSummaryBufferMemory

from langchain.schema import format_document


# document loaders
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    DirectoryLoader,
    CSVLoader,
    Docx2txtLoader,
)

# text_splitter
from langchain.text_splitter import (
    RecursiveCharacterTextSplitter,
    CharacterTextSplitter,
)

# OutputParser
from langchain_core.output_parsers import StrOutputParser

# Import chroma as the vector store
from langchain_community.vectorstores import Chroma

# Contextual_compression
from langchain.retrievers.document_compressors import DocumentCompressorPipeline
from langchain_community.document_transformers import (
    EmbeddingsRedundantFilter,
    LongContextReorder,
)
from langchain.retrievers.document_compressors import EmbeddingsFilter
from langchain.retrievers import ContextualCompressionRetriever

# Cohere
from langchain.retrievers.document_compressors import CohereRerank
from langchain_community.llms import Cohere

# HuggingFace
from langchain_community.embeddings import HuggingFaceInferenceAPIEmbeddings
from langchain_community.llms import HuggingFaceHub

# Import streamlit
import streamlit as st

####################################################################
#              Config: LLM services, assistant language,...
####################################################################
list_LLM_providers = [
    ":rainbow[**OpenAI**]",
    "**Google Generative AI**",
    ":hugging_face: **HuggingFace**",
]

dict_welcome_message = {
    "english": "How can I assist you today?",
    "french": "Comment puis-je vous aider aujourd'hui ?",
    "spanish": "¿Cómo puedo ayudarle hoy?",
    "german": "Wie kann ich Ihnen heute helfen?",
    "russian": "Чем я могу помочь вам сегодня?",
    "chinese": "我今天能帮你什么？",
    "arabic": "كيف يمكنني مساعدتك اليوم؟",
    "portuguese": "Como posso ajudá-lo hoje?",
    "italian": "Come posso assistervi oggi?",
    "Japanese": "今日はどのようなご用件でしょうか?",
}

list_retriever_types = [
    "Cohere reranker",
    "Contextual compression",
    "Vectorstore backed retriever",
]

TMP_DIR = Path(__file__).resolve().parent.joinpath("data", "tmp")
LOCAL_VECTOR_STORE_DIR = (
    Path(__file__).resolve().parent.joinpath("data", "vector_stores")
)

####################################################################
#            Create app interface with streamlit
####################################################################
st.set_page_config(page_title="Chat With Your Data")

st.title("🤖 RAG chatbot")

# API keys
st.session_state.openai_api_key = ""
st.session_state.google_api_key = ""
st.session_state.cohere_api_key = ""
st.session_state.hf_api_key = ""


def expander_model_parameters(
    LLM_provider="OpenAI",
    text_input_API_key="OpenAI API Key - [Get an API key](https://platform.openai.com/account/api-keys)",
    list_models=["gpt-3.5-turbo-0125", "gpt-3.5-turbo", "gpt-4-turbo-preview"],
):
    """Add a text_input (for API key) and a streamlit expander containing models and parameters."""
    st.session_state.LLM_provider = LLM_provider

    if LLM_provider == "OpenAI":
        st.session_state.openai_api_key = st.text_input(
            text_input_API_key,
            type="password",
            placeholder="insert your API key",
        )
        st.session_state.google_api_key = ""
        st.session_state.hf_api_key = ""

    if LLM_provider == "Google":
        st.session_state.google_api_key = st.text_input(
            text_input_API_key,
            type="password",
            placeholder="insert your API key",
        )
        st.session_state.openai_api_key = ""
        st.session_state.hf_api_key = ""

    if LLM_provider == "HuggingFace":
        st.session_state.hf_api_key = st.text_input(
            text_input_API_key,
            type="password",
            placeholder="insert your API key",
        )
        st.session_state.openai_api_key = ""
        st.session_state.google_api_key = ""

    with st.expander("**Models and parameters**"):
        st.session_state.selected_model = st.selectbox(
            f"Choose {LLM_provider} model", list_models
        )

        # model parameters
        st.session_state.temperature = st.slider(
            "temperature",
            min_value=0.0,
            max_value=1.0,
            value=0.5,
            step=0.1,
        )
        st.session_state.top_p = st.slider(
            "top_p",
            min_value=0.0,
            max_value=1.0,
            value=0.95,
            step=0.05,
        )


def sidebar_and_documentChooser():
    """Create the sidebar and the a tabbed pane: the first tab contains a document chooser (create a new vectorstore);
    the second contains a vectorstore chooser (open an old vectorstore)."""

    with st.sidebar:
        st.caption(
            "🚀 A retrieval augmented generation chatbot powered by 🔗 Langchain, Cohere, OpenAI, Google Generative AI and 🤗"
        )
        st.write("")

        llm_chooser = st.radio(
            "Select provider",
            list_LLM_providers,
            captions=[
                "[OpenAI pricing page](https://openai.com/pricing)",
                "Rate limit: 60 requests per minute.",
                "**Free access.**",
            ],
        )

        st.divider()
        if llm_chooser == list_LLM_providers[0]:
            expander_model_parameters(
                LLM_provider="OpenAI",
                text_input_API_key="OpenAI API Key - [Get an API key](https://platform.openai.com/account/api-keys)",
                list_models=[
                    "gpt-3.5-turbo-0125",
                    "gpt-3.5-turbo",
                    "gpt-4-turbo-preview",
                ],
            )

        if llm_chooser == list_LLM_providers[1]:
            expander_model_parameters(
                LLM_provider="Google",
                text_input_API_key="Google API Key - [Get an API key](https://makersuite.google.com/app/apikey)",
                list_models=["gemini-pro"],
            )
        if llm_chooser == list_LLM_providers[2]:
            expander_model_parameters(
                LLM_provider="HuggingFace",
                text_input_API_key="HuggingFace API key - [Get an API key](https://huggingface.co/settings/tokens)",
                list_models=["mistralai/Mistral-7B-Instruct-v0.2"],
            )
        # Assistant language
        st.write("")
        st.session_state.assistant_language = st.selectbox(
            f"Assistant language", list(dict_welcome_message.keys())
        )

        st.divider()
        st.subheader("Retrievers")
        retrievers = list_retriever_types
        if st.session_state.selected_model == "gpt-3.5-turbo":
            # for "gpt-3.5-turbo", we will not use the vectorstore backed retriever
            # there is a high risk of exceeding the max tokens limit (4096).
            retrievers = list_retriever_types[:-1]

        st.session_state.retriever_type = st.selectbox(
            f"Select retriever type", retrievers
        )
        st.write("")
        if st.session_state.retriever_type == list_retriever_types[0]:  # Cohere
            st.session_state.cohere_api_key = st.text_input(
                "Coher API Key - [Get an API key](https://dashboard.cohere.com/api-keys)",
                type="password",
                placeholder="insert your API key",
            )

        st.write("\n\n")
        st.write(
            f"ℹ _Your {st.session_state.LLM_provider} API key, '{st.session_state.selected_model}' parameters, \
            and {st.session_state.retriever_type} are only considered when loading or creating a vectorstore._"
        )

    # Tabbed Pane: Create a new Vectorstore | Open a saved Vectorstore

    tab_new_vectorstore, tab_open_vectorstore = st.tabs(
        ["Create a new Vectorstore", "Open a saved Vectorstore"]
    )
    with tab_new_vectorstore:
        # 1. Select documnets
        st.session_state.uploaded_file_list = st.file_uploader(
            label="**Select documents**",
            accept_multiple_files=True,
            type=(["pdf", "txt", "docx", "csv"]),
        )
        # 2. Process documents
        st.session_state.vector_store_name = st.text_input(
            label="**Documents will be loaded, embedded and ingested into a vectorstore (Chroma dB). Please provide a valid dB name.**",
            placeholder="Vectorstore name",
        )
        # 3. Add a button to process documnets and create a Chroma vectorstore

        st.button("Create Vectorstore", on_click=chain_RAG_blocks)
        try:
            if st.session_state.error_message != "":
                st.warning(st.session_state.error_message)
        except:
            pass

    with tab_open_vectorstore:
        # Open a saved Vectorstore
        st.write("Please select a Vectorstore:")
        
        # Replace Tkinter file dialog with Streamlit's file_uploader
        st.session_state.selected_vectorstore_name = ""
        
        # Use a text input for the vectorstore path instead of Tkinter dialog
        vectorstore_path = st.text_input(
            "Enter the path to your vectorstore directory:",
            placeholder="e.g., /path/to/your/vectorstore"
        )
        
        if st.button("Load Vectorstore"):
            # Check inputs
            error_messages = []
            if (
                not st.session_state.openai_api_key
                and not st.session_state.google_api_key
                and not st.session_state.hf_api_key
            ):
                error_messages.append(
                    f"insert your {st.session_state.LLM_provider} API key"
                )

            if (
                st.session_state.retriever_type == list_retriever_types[0]
                and not st.session_state.cohere_api_key
            ):
                error_messages.append(f"insert your Cohere API key")

            if len(error_messages) == 1:
                st.session_state.error_message = "Please " + error_messages[0] + "."
                st.warning(st.session_state.error_message)
            elif len(error_messages) > 1:
                st.session_state.error_message = (
                    "Please "
                    + ", ".join(error_messages[:-1])
                    + ", and "
                    + error_messages[-1]
                    + "."
                )
                st.warning(st.session_state.error_message)

            # if API keys are inserted, start loading Chroma index, then create retriever and ConversationalRetrievalChain
            elif vectorstore_path == "":
                st.info("Please enter a valid path.")

            else:
                with st.spinner("Loading vectorstore..."):
                    st.session_state.selected_vectorstore_name = (
                        vectorstore_path.split("/")[-1]
                    )
                    try:
                        # 1. load Chroma vectorestore
                        embeddings = select_embeddings_model()
                        st.session_state.vector_store = Chroma(
                            embedding_function=embeddings,
                            persist_directory=vectorstore_path,
                        )

                        # 2. create retriever
                        st.session_state.retriever = create_retriever(
                            vector_store=st.session_state.vector_store,
                            embeddings=embeddings,
                            retriever_type=st.session_state.retriever_type,
                            base_retriever_search_type="similarity",
                            base_retriever_k=16,
                            compression_retriever_k=20,
                            cohere_api_key=st.session_state.cohere_api_key,
                            cohere_model="rerank-multilingual-v2.0",
                            cohere_top_n=10,
                        )

                        # 3. create memory and ConversationalRetrievalChain
                        (
                            st.session_state.chain,
                            st.session_state.memory,
                        ) = create_ConversationalRetrievalChain(
                            retriever=st.session_state.retriever,
                            chain_type="stuff",
                            language=st.session_state.assistant_language,
                        )

                        # 4. clear chat_history
                        clear_chat_history()

                        st.info(
                            f"**{st.session_state.selected_vectorstore_name}** is loaded successfully."
                        )

                    except Exception as e:
                        st.error(e)


####################################################################
#        Process documents and create vectorstor (Chroma dB)
####################################################################
def delte_temp_files():
    """delete files from the './data/tmp' folder"""
    files = glob.glob(TMP_DIR.as_posix() + "/*")
    for f in files:
        try:
            os.remove(f)
        except:
            pass


def langchain_document_loader():
    """
    Crete documnet loaders for PDF, TXT and CSV files.
    https://python.langchain.com/docs/modules/data_connection/document_loaders/file_directory
    """

    documents = []

    txt_loader = DirectoryLoader(
        TMP_DIR.as_posix(), glob="**/*.txt", loader_cls=TextLoader, show_progress=True
    )
    documents.extend(txt_loader.load())

    pdf_loader = DirectoryLoader(
        TMP_DIR.as_posix(), glob="**/*.pdf", loader_cls=PyPDFLoader, show_progress=True
    )
    documents.extend(pdf_loader.load())

    csv_loader = DirectoryLoader(
        TMP_DIR.as_posix(), glob="**/*.csv", loader_cls=CSVLoader, show_progress=True,
        loader_kwargs={"encoding":"utf8"}
    )
    documents.extend(csv_loader.load())

    doc_loader = DirectoryLoader(
        TMP_DIR.as_posix(),
        glob="**/*.docx",
        loader_cls=Docx2txtLoader,
        show_progress=True,
    )
    documents.extend(doc_loader.load())
    return documents


def split_documents_to_chunks(documents):
    """Split documents to chunks using RecursiveCharacterTextSplitter."""

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1600, chunk_overlap=200)
    chunks = text_splitter.split_documents(documents)
    return chunks


def select_embeddings_model():
    """Select embeddings models: OpenAIEmbeddings or GoogleGenerativeAIEmbeddings."""
    if st.session_state.LLM_provider == "OpenAI":
        embeddings = OpenAIEmbeddings(api_key=st.session_state.openai_api_key)

    if st.session_state.LLM_provider == "Google":
        embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001", google_api_key=st.session_state.google_api_key
        )

    if st.session_state.LLM_provider == "HuggingFace":
        embeddings = HuggingFaceInferenceAPIEmbeddings(
            api_key=st.session_state.hf_api_key, model_name="thenlper/gte-large"
        )

    return embeddings


def create_retriever(
    vector_store,
    embeddings,
    retriever_type="Contextual compression",
    base_retriever_search_type="semilarity",
    base_retriever_k=16,
    compression_retriever_k=20,
    cohere_api_key="",
    cohere_model="rerank-multilingual-v2.0",
    cohere_top_n=10,
):
    """
    create a retriever which can be a:
        - Vectorstore backed retriever: this is the base retriever.
        - Contextual compression retriever: We wrap the the base retriever in a ContextualCompressionRetriever.
            The compressor here is a Document Compressor Pipeline, which splits documents
            to smaller chunks, removes redundant documents, filters the top relevant documents,
            and reorder the documents so that the most relevant are at beginning / end of the list.
        - Cohere_reranker: CohereRerank endpoint is used to reorder the results based on relevance.

    Parameters:
        vector_store: Chroma vector database.
        embeddings: OpenAIEmbeddings or GoogleGenerativeAIEmbeddings.

        retriever_type (str): in [Vectorstore backed retriever,Contextual compression,Cohere reranker]. default = Cohere reranker

        base_retreiver_search_type: search_type in ["similarity", "mmr", "similarity_score_threshold"], default = similarity.
        base_retreiver_k: The most similar vectors are returned (default k = 16).

        compression_retriever_k: top k documents returned by the compression retriever, default = 20

        cohere_api_key: Cohere API key
        cohere_model (str): model used by Cohere, in ["rerank-multilingual-v2.0","rerank-english-v2.0"]
        cohere_top_n: top n documents returned bu Cohere, default = 10

    """

    base_retriever = Vectorstore_backed_retriever(
        vectorstore=vector_store,
        search_type=base_retriever_search_type,
        k=base_retriever_k,
        score_threshold=None,
    )

    if retriever_type == "Vectorstore backed retriever":
        return base_retriever

    elif retriever_type == "Contextual compression":
        compression_retriever = create_compression_retriever(
            embeddings=embeddings,
            base_retriever=base_retriever,
            k=compression_retriever_k,
        )
        return compression_retriever

    elif retriever_type == "Cohere reranker":
        cohere_retriever = CohereRerank_retriever(
            base_retriever=base_retriever,
            cohere_api_key=cohere_api_key,
            cohere_model=cohere_model,
            top_n=cohere_top_n,
        )
        return cohere_retriever
    else:
        pass


def Vectorstore_backed_retriever(
    vectorstore, search_type="similarity", k=4, score_threshold=None
):
    """create a vectorsore-backed retriever
    Parameters:
        search_type: Defines the type of search that the Retriever should perform.
            Can be "similarity" (default), "mmr", or "similarity_score_threshold"
        k: number of documents to return (Default: 4)
        score_threshold: Minimum relevance threshold for similarity_score_threshold (default=None)
    """
    search_kwargs = {}
    if k is not None:
        search_kwargs["k"] = k
    if score_threshold is not None:
        search_kwargs["score_threshold"] = score_threshold

    retriever = vectorstore.as_retriever(
        search_type=search_type, search_kwargs=search_kwargs
    )
    return retriever


def create_compression_retriever(
    embeddings, base_retriever, chunk_size=500, k=16, similarity_threshold=None
):
    """Build a ContextualCompressionRetriever.
    We wrap the the base_retriever (a Vectorstore-backed retriever) in a ContextualCompressionRetriever.
    The compressor here is a Document Compressor Pipeline, which splits documents
    to smaller chunks, removes redundant documents, filters the top relevant documents,
    and reorder the documents so that the most relevant are at beginning / end of the list.

    Parameters:
        embeddings: OpenAIEmbeddings or GoogleGenerativeAIEmbeddings.
        base_retriever: a Vectorstore-backed retriever.
        chunk_size (int): Docs will be splitted into smaller chunks using a CharacterTextSplitter with a default chunk_size of 500.
        k (int): top k relevant documents to the query are filtered using the EmbeddingsFilter. default =16.
        similarity_threshold : similarity_threshold of the  EmbeddingsFilter. default =None
    """

    # 1. splitting docs into smaller chunks
    splitter = CharacterTextSplitter(
        chunk_size=chunk_size, chunk_overlap=0, separator=". "
    )

    # 2. removing redundant documents
    redundant_filter = EmbeddingsRedundantFilter(embeddings=embeddings)

    # 3. filtering based on relevance to the query
    relevant_filter = EmbeddingsFilter(
        embeddings=embeddings, k=k, similarity_threshold=similarity_threshold
    )

    # 4. Reorder the documents

    # Less relevant document will be at the middle of the list and more relevant elements at beginning / end.
    # Reference: https://python.langchain.com/docs/modules/data_connection/retrievers/long_context_reorder
    reordering = LongContextReorder()

    # 5. create compressor pipeline and retriever
    pipeline_compressor = DocumentCompressorPipeline(
        transformers=[splitter, redundant_filter, relevant_filter, reordering]
    )
    compression_retriever = ContextualCompressionRetriever(
        base_compressor=pipeline_compressor, base_retriever=base_retriever
    )

    return compression_retriever


def CohereRerank_retriever(
    base_retriever, cohere_api_key, cohere_model="rerank-multilingual-v2.0", top_n=10
):
    """Build a ContextualCompressionRetriever using CohereRerank endpoint to reorder the results
    based on relevance to the query.

    Parameters:
       base_retriever: a Vectorstore-backed retriever
       cohere_api_key: the Cohere API key
       cohere_model: the Cohere model, in ["rerank-multilingual-v2.0","rerank-english-v2.0"], default = "rerank-multilingual-v2.0"
       top_n: top n results returned by Cohere rerank. default = 10.
    """

    compressor = CohereRerank(
        cohere_api_key=cohere_api_key, model=cohere_model, top_n=top_n
    )

    retriever_Cohere = ContextualCompressionRetriever(
        base_compressor=compressor, base_retriever=base_retriever
    )
    return retriever_Cohere


def chain_RAG_blocks():
    """The RAG system is composed of:
    - 1. Retrieval: includes document loaders, text splitter, vectorstore and retriever.
    - 2. Memory.
    - 3. Converstaional Retreival chain.
    """
    with st.spinner("Creating vectorstore..."):
        # Check inputs
        error_messages = []
        if (
            not st.session_state.openai_api_key
            and not st.session_state.google_api_key
            and not st.session_state.hf_api_key
        ):
            error_messages.append(
                f"insert your {st.session_state.LLM_provider} API key"
            )

        if (
            st.session_state.retriever_type == list_retriever_types[0]
            and not st.session_state.cohere_api_key
        ):
            error_messages.append(f"insert your Cohere API key")
        if not st.session_state.uploaded_file_list:
            error_messages.append("select documents to upload")
        if st.session_state.vector_store_name == "":
            error_messages.append("provide a Vectorstore name")

        if len(error_messages) == 1:
            st.session_state.error_message = "Please " + error_messages[0] + "."
        elif len(error_messages) > 1:
            st.session_state.error_message = (
                "Please "
                + ", ".join(error_messages[:-1])
                + ", and "
                + error_messages[-1]
                + "."
            )
        else:
            st.session_state.error_message = ""
            try:
                # 1. Delete old temp files
                delte_temp_files()

                # 2. Upload selected documents to temp directory
                if st.session_state.uploaded_file_list is not None:
                    for uploaded_file in st.session_state.uploaded_file_list:
                        error_message = ""
                        try:
                            temp_file_path = os.path.join(
                                TMP_DIR.as_posix(), uploaded_file.name
                            )
                            with open(temp_file_path, "wb") as temp_file:
                                temp_file.write(uploaded_file.read())
                        except Exception as e:
                            error_message += str(e)
                    if error_message != "":
                        st.warning(f"Errors: {error_message}")

                    # 3. Load documents with Langchain loaders
                    documents = langchain_document_loader()

                    # 4. Split documents to chunks
                    chunks = split_documents_to_chunks(documents)
                    # 5. Embeddings
                    embeddings = select_embeddings_model()

                    # 6. Create a vectorstore
                    persist_directory = (
                        LOCAL_VECTOR_STORE_DIR.as_posix()
                        + "/"
                        + st.session_state.vector_store_name
                    )

                    try:
                        st.session_state.vector_store = Chroma.from_documents(
                            documents=chunks,
                            embedding=embeddings,
                            persist_directory=persist_directory,
                        )
                        st.info(
                            f"Vectorstore **{st.session_state.vector_store_name}** is created succussfully."
                        )

                        # 7. Create retriever
                        st.session_state.retriever = create_retriever(
                            vector_store=st.session_state.vector_store,
                            embeddings=embeddings,
                            retriever_type=st.session_state.retriever_type,
                            base_retriever_search_type="similarity",
                            base_retriever_k=16,
                            compression_retriever_k=20,
                            cohere_api_key=st.session_state.cohere_api_key,
                            cohere_model="rerank-multilingual-v2.0",
                            cohere_top_n=10,
                        )

                        # 8. Create memory and ConversationalRetrievalChain
                        (
                            st.session_state.chain,
                            st.session_state.memory,
                        ) = create_ConversationalRetrievalChain(
                            retriever=st.session_state.retriever,
                            chain_type="stuff",
                            language=st.session_state.assistant_language,
                        )

                        # 9. Cclear chat_history
                        clear_chat_history()

                    except Exception as e:
                        st.error(e)

            except Exception as error:
                st.error(f"An error occurred: {error}")


####################################################################
#                       Create memory
####################################################################


def create_memory(model_name="gpt-3.5-turbo", memory_max_token=None):
    """Creates a ConversationSummaryBufferMemory for gpt-3.5-turbo
    Creates a ConversationBufferMemory for the other models"""

    if model_name == "gpt-3.5-turbo":
        if memory_max_token is None:
            memory_max_token = 1024  # max_tokens for 'gpt-3.5-turbo' = 4096
        memory = ConversationSummaryBufferMemory(
            max_token_limit=memory_max_token,
            llm=ChatOpenAI(
                model_name="gpt-3.5-turbo",
                openai_api_key=st.session_state.openai_api_key,
                temperature=0.1,
            ),
            return_messages=True,
            memory_key="chat_history",
            output_key="answer",
            input_key="question",
        )
    else:
        memory = ConversationBufferMemory(
            return_messages=True,
            memory_key="chat_history",
            output_key="answer",
            input_key="question",
        )
    return memory


####################################################################
#          Create ConversationalRetrievalChain with memory
####################################################################


def answer_template(language="english"):
    """Pass the standalone question along with the chat history and context
    to the `LLM` wihch will answer."""

    template = f"""Answer the question at the end, using only the following context (delimited by <context></context>).
Your answer must be in the language at the end. 

<context>
{{chat_history}}

{{context}} 
</context>

Question: {{question}}

Language: {language}.
"""
    return template


def create_ConversationalRetrievalChain(
    retriever,
    chain_type="stuff",
    language="english",
):
    """Create a ConversationalRetrievalChain.
    First, it passes the follow-up question along with the chat history to an LLM which rephrases
    the question and generates a standalone query.
    This query is then sent to the retriever, which fetches relevant documents (context)
    and passes them along with the standalone question and chat history to an LLM to answer.
    """

    # 1. Define the standalone_question prompt.
    # Pass the follow-up question along with the chat history to the `condense_question_llm`
    # which rephrases the question and generates a standalone question.

    condense_question_prompt = PromptTemplate(
        input_variables=["chat_history", "question"],
        template="""Given the following conversation and a follow up question, 
rephrase the follow up question to be a standalone question, in its original language.\n\n
Chat History:\n{chat_history}\n
Follow Up Input: {question}\n
Standalone question:""",
    )

    # 2. Define the answer_prompt
    # Pass the standalone question + the chat history + the context (retrieved documents)
    # to the `LLM` wihch will answer

    answer_prompt = ChatPromptTemplate.from_template(answer_template(language=language))

    # 3. Add ConversationSummaryBufferMemory for gpt-3.5, and ConversationBufferMemory for the other models
    memory = create_memory(st.session_state.selected_model)

    # 4. Instantiate LLMs: standalone_query_generation_llm & response_generation_llm
    if st.session_state.LLM_provider == "OpenAI":
        standalone_query_generation_llm = ChatOpenAI(
            api_key=st.session_state.openai_api_key,
            model=st.session_state.selected_model,
            temperature=0.1,
        )
        response_generation_llm = ChatOpenAI(
            api_key=st.session_state.openai_api_key,
            model=st.session_state.selected_model,
            temperature=st.session_state.temperature,
            model_kwargs={"top_p": st.session_state.top_p},
        )
    if st.session_state.LLM_provider == "Google":
        standalone_query_generation_llm = ChatGoogleGenerativeAI(
            google_api_key=st.session_state.google_api_key,
            model=st.session_state.selected_model,
            temperature=0.1,
            convert_system_message_to_human=True,
        )
        response_generation_llm = ChatGoogleGenerativeAI(
            google_api_key=st.session_state.google_api_key,
            model=st.session_state.selected_model,
            temperature=st.session_state.temperature,
            top_p=st.session_state.top_p,
            convert_system_message_to_human=True,
        )

    if st.session_state.LLM_provider == "HuggingFace":
        standalone_query_generation_llm = HuggingFaceHub(
            repo_id=st.session_state.selected_model,
            huggingfacehub_api_token=st.session_state.hf_api_key,
            model_kwargs={
                "temperature": 0.1,
                "top_p": 0.95,
                "do_sample": True,
                "max_new_tokens": 1024,
            },
        )
        response_generation_llm = HuggingFaceHub(
            repo_id=st.session_state.selected_model,
            huggingfacehub_api_token=st.session_state.hf_api_key,
            model_kwargs={
                "temperature": st.session_state.temperature,
                "top_p": st.session_state.top_p,
                "do_sample": True,
                "max_new_tokens": 1024,
            },
        )

    # 5. Create the ConversationalRetrievalChain

    chain = ConversationalRetrievalChain.from_llm(
        condense_question_prompt=condense_question_prompt,
        combine_docs_chain_kwargs={"prompt": answer_prompt},
        condense_question_llm=standalone_query_generation_llm,
        llm=response_generation_llm,
        memory=memory,
        retriever=retriever,
        chain_type=chain_type,
        verbose=False,
        return_source_documents=True,
    )

    return chain, memory


def clear_chat_history():
    """clear chat history and memory."""
    # 1. re-initialize messages
    st.session_state.messages = [
        {
            "role": "assistant",
            "content": dict_welcome_message[st.session_state.assistant_language],
        }
    ]
    # 2. Clear memory (history)
    try:
        st.session_state.memory.clear()
    except:
        pass

    



from evaluation import ragas_metrics, consistency_llm
from ragas import evaluate
import nltk, numpy as np

LOW_SENTENCE = .25   # highlight thresholds
MID_SENTENCE = .50

# ---------- new helper ----------
def highlight_with_metrics(text: str, metrics: dict) -> str:
    """
    Color-code text based on multiple metrics:
    - Red: Any metric < 0.25 (low confidence/quality)
    - Orange: Any metric < 0.50 (medium confidence/quality)
    - No highlight (black text): All metrics >= 0.50 (high confidence/quality)
    """
    sentences = nltk.sent_tokenize(text)
    html = ""
    
    # Get metrics
    answer_relevancy = metrics.get('answer_relevancy', 0.0)
    faithfulness = metrics.get('faithfulness', 0.0)
    context_utilization = metrics.get('context_utilization', 0.0)
    consistency = metrics.get('consistency', 0.0)
    
    for sent in sentences:
        # Determine which metrics are concerning for this sentence
        low_metrics = []
        if answer_relevancy < 0.95:
            low_metrics.append(f"Relevancy: {answer_relevancy:.2f}")
        if faithfulness < .95:
            low_metrics.append(f"Faithfulness: {faithfulness:.2f}")
        if context_utilization < 0.50:
            low_metrics.append(f"Context Use: {context_utilization:.2f}")
        if consistency < 0.50:
            low_metrics.append(f"Consistency: {consistency:.2f}")
            
        if low_metrics:
            if any(score < 0.25 for score in [answer_relevancy, faithfulness, context_utilization, consistency]):
                color = "#FF4B4B"  # bright red
                bg_color = "#FFE5E5"  # light red background
                confidence = "Low confidence"
            else:
                color = "#FFA500"  # orange
                bg_color = "#FFF3E0"  # light orange background
                confidence = "Medium confidence"
                
            tooltip = f"""Confidence: {confidence}
Concerns: {', '.join(low_metrics)}"""
            
            html += f'<span style="color: {color}; background-color: {bg_color}; padding: 2px 4px; border-radius: 3px; margin: 0 2px;" title="{tooltip}">{sent}</span> '
        else:
            # No highlighting needed - use black text
            html += f'<span style="color: black;">{sent}</span> '
            
    return html


def get_response_from_LLM(prompt):
    """invoke the LLM, get response, and display results with evaluation metrics."""
    try:
        # Check for OpenAI API key
        if not st.session_state.get('openai_api_key'):
            st.warning("Please provide your OpenAI API key to enable evaluation metrics.")
            return

        # 1. Invoke LLM
        response = st.session_state.chain.invoke({"question": prompt})
        answer = response["answer"]
        docs = response["source_documents"]

        # 2. Display results
        st.session_state.messages.append({"role": "user", "content": prompt})
        st.session_state.messages.append({"role": "assistant", "content": answer})
        st.chat_message("user").write(prompt)
        
        with st.chat_message("assistant"):
            try:
                # Get RAGAS scores and consistency score
                os.environ["OPENAI_API_KEY"] = st.session_state.openai_api_key
                eval_data = {
                    "question": [prompt],
                    "answer": [answer],
                    "contexts": [[d.page_content for d in docs]]
                }
                
                from datasets import Dataset
                dataset = Dataset.from_dict(eval_data)
                ragas_scores = evaluate(dataset, metrics=ragas_metrics)
                
                # Get consistency score
                consistency_score = consistency_llm.evaluate_strings(
                    prediction=answer, 
                    reference=""
                ).score / 5.0  # Normalize to 0-1
                
                # Combine all scores
                all_scores = {**ragas_scores}
                for metric_name, score in all_scores.items():
                    if isinstance(score, (int, float)):
                        all_scores[metric_name] = score
                all_scores["consistency"] = consistency_score
                
                # Display highlighted answer with all metrics considered
                st.markdown(highlight_with_metrics(answer, all_scores), unsafe_allow_html=True)

                # Show warnings for any metric below 0.5
                low_metrics = []
                for metric_name, score in all_scores.items():
                    if isinstance(score, (int, float)) and score < 0.5:
                        low_metrics.append(f"**{metric_name}** ({score:.2f})")
                
                if low_metrics:
                    warning_msg = "⚠️ Quality concerns detected in: " + ", ".join(low_metrics)
                    st.warning(warning_msg + "\n\n" + """
                    • Red highlighting: Major concerns (scores < 0.25)
                    • Orange highlighting: Some concerns (scores < 0.50)
                    • Green highlighting: Good quality (scores ≥ 0.50)
                    
                    Hover over highlighted text to see detailed scores.""")

                # Add buttons for adding content to Word document
                col1, col2 = st.columns([1, 1])
                with col1:
                    if st.button("📝 Add Q&A to Report", key=f"add_qa_{len(st.session_state.messages)}"):
                        content = {
                            'question': prompt,
                            'answer': answer,
                            'metrics': all_scores,
                            'sources': [{
                                'source': doc.metadata.get('source', 'Unknown'),
                                'page': doc.metadata.get('page', ''),
                                'content': doc.page_content
                            } for doc in docs]
                        }
                        st.session_state.selected_content.append(content)
                        st.success("Added to report!")
                
                with col2:
                    if st.button("📝 Add Q&A Only", key=f"add_qa_only_{len(st.session_state.messages)}"):
                        content = {
                            'question': prompt,
                            'answer': answer
                        }
                        st.session_state.selected_content.append(content)
                        st.success("Added to report (Q&A only)!")

            except Exception as e:
                # Fallback to plain answer if highlighting fails
                st.markdown(answer)
                st.warning(f"Could not evaluate response: {str(e)}")

            # Display evaluation metrics
            with st.expander("📊 Quality metrics"):
                try:
                    if all_scores:
                        st.markdown("""
                        **Understanding the metrics:**
                        - **Answer relevancy**: How well the answer addresses the question
                        - **Faithfulness**: How accurately it reflects the source documents
                        - **Context utilization**: How effectively it uses the provided context
                        - **Consistency**: How internally consistent the answer is
                        """)
                        st.table({k: f"{v:.2f}" for k, v in all_scores.items()})
                    else:
                        st.info("No evaluation metrics available")

                except Exception as e:
                    st.warning(f"Could not display metrics: {str(e)}")

            # Display source documents
            with st.expander("📑 Source documents"):
                for doc in docs:
                    try:
                        page = " (Page: " + str(doc.metadata["page"]) + ")"
                    except:
                        page = ""
                    st.markdown(
                        f"**Source: {doc.metadata['source']}{page}**\n\n{doc.page_content[:800]}..."
                    )

    except Exception as e:
        st.warning(str(e))


####################################################################
#                         Chatbot
####################################################################
def chatbot():
    # Initialize session state for document handling
    initialize_session_state()
    
    sidebar_and_documentChooser()
    
    # Add Word document controls in sidebar
    with st.sidebar:
        st.divider()
        st.subheader("📝 Report Controls")
        
        # Show number of selected items
        num_selected = len(st.session_state.selected_content)
        st.write(f"Selected items: {num_selected}")
        
        col1, col2 = st.columns(2)
        
        # Clear selection button
        if col1.button("Clear Selection") and num_selected > 0:
            st.session_state.selected_content = []
            st.success("Selection cleared!")
            
        # Compile and download button
        if col2.button("Compile Report") and num_selected > 0:
            save_path, error = save_compiled_document()
            if save_path:
                with open(save_path, "rb") as file:
                    st.download_button(
                        label="Download Report",
                        data=file,
                        file_name=os.path.basename(save_path),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.error(f"Error saving document: {error}")
    
    st.divider()
    col1, col2 = st.columns([7, 3])
    with col1:
        st.subheader("Chat with your data")
    with col2:
        st.button("Clear Chat History", on_click=clear_chat_history)

    # Initialize messages if not exists
    if "messages" not in st.session_state:
        st.session_state["messages"] = [
            {
                "role": "assistant",
                "content": dict_welcome_message[st.session_state.assistant_language],
            }
        ]

    # Display existing messages
    for msg in st.session_state.messages:
        st.chat_message(msg["role"]).write(msg["content"])

    # Check for API keys and chain initialization
    if prompt := st.chat_input():
        if (
            not st.session_state.openai_api_key
            and not st.session_state.google_api_key
            and not st.session_state.hf_api_key
        ):
            st.info(
                f"Please insert your {st.session_state.LLM_provider} API key to continue."
            )
            st.stop()
        
        # Check if chain is initialized
        if not hasattr(st.session_state, 'chain') or st.session_state.chain is None:
            st.warning("Please load or create a vectorstore first before chatting.")
            st.stop()
            
        with st.spinner("Running..."):
            get_response_from_LLM(prompt=prompt)


def initialize_session_state():
    """Initialize session state variables for document handling"""
    if 'selected_content' not in st.session_state:
        st.session_state.selected_content = []
    if 'current_doc' not in st.session_state:
        st.session_state.current_doc = None

def create_word_document():
    """Create a new Word document with proper formatting"""
    doc = Document()
    # Add title
    title = doc.add_heading('Chat History Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    timestamp = doc.add_paragraph()
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timestamp.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    doc.add_paragraph()  # Add some space
    return doc

def add_qa_to_word(doc, question, answer, metrics=None, sources=None):
    """Add a Q&A exchange with optional metrics and sources to the Word document"""
    # Add question
    q_para = doc.add_paragraph()
    q_para.add_run("Question: ").bold = True
    q_para.add_run(question)
    
    # Add answer
    a_para = doc.add_paragraph()
    a_para.add_run("Answer: ").bold = True
    a_para.add_run(html.unescape(answer))
    
    # Add metrics if provided
    if metrics:
        doc.add_heading("Quality Metrics", level=2)
        metrics_table = doc.add_table(rows=1, cols=2)
        metrics_table.style = 'Table Grid'
        header_cells = metrics_table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Score'
        
        for metric, value in metrics.items():
            row_cells = metrics_table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = f"{value:.2f}"
    
    # Add sources if provided
    if sources:
        doc.add_heading("Source Documents", level=2)
        for source in sources:
            s_para = doc.add_paragraph()
            s_para.add_run(f"Source: {source['source']}")
            if source.get('page'):
                s_para.add_run(f" (Page: {source['page']})")
            s_para.add_run(f"\n{source['content'][:800]}...")
    
    doc.add_paragraph()  # Add space between entries
    doc.add_paragraph("=" * 50)  # Add separator
    doc.add_paragraph()

def save_compiled_document():
    """Save the compiled Word document with all selected content"""
    if not st.session_state.selected_content:
        return None, "No content selected to save"
    
    try:
        doc = create_word_document()
        
        for content in st.session_state.selected_content:
            add_qa_to_word(
                doc,
                content['question'],
                content['answer'],
                content.get('metrics'),
                content.get('sources')
            )
        
        # Save the document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"chat_history_{timestamp}.docx"
        save_path = os.path.join(os.path.dirname(__file__), "exports", filename)
        
        # Create exports directory if it doesn't exist
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        
        doc.save(save_path)
        return save_path, None
    except Exception as e:
        return None, str(e)


if __name__ == "__main__":
    chatbot()
